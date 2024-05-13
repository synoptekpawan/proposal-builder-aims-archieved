"""
This code uses python==3.11.4, and open_ai==0.28.1, langchain==0.0.354
"""
# load packages
import streamlit as st
from streamlit_quill import st_quill
from langchain.prompts import PromptTemplate
# from langchain_community.llms import CTransformers, OpenAI
import openai
import requests
import pypandoc
from docx import Document
from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.text import WD_BREAK
# from azure.storage.blob import BlobServiceClient
from datetime import datetime
from docxtpl import DocxTemplate, RichText
import pandas as pd
import jinja2
import os
import re
import io
import time
import asyncio
from dotenv import load_dotenv,find_dotenv
load_dotenv(find_dotenv())

from prompts.allDoc_Prompts4 import *
# from azure.identity import DefaultAzureCredential, ManagedIdentityCredential
# from azure.keyvault.secrets import SecretClient

# local folder setup
if not os.path.exists(r'./generated_docs/'):
    os.mkdir(r'./generated_docs/')

# azure open ai setup
openai.api_type = "azure"
openai.api_base = "https://dev-aims-ai.openai.azure.com/"
openai.api_version = "2023-12-01-preview"
openai.api_key = os.getenv("OPENAI_API_KEY_AZURE")

deployment_id ="gpt-35-aims" # "gpt4-aims"#Add your deployment ID here
engine = "gpt-35-instruct"

# Azure AI Search setup
search_endpoint = "https://cogs-aims.search.windows.net"; # Add your Azure AI Search endpoint here
search_key = os.getenv("SEARCH_KEY") # Add your Azure AI Search admin key here
search_index_name = "try-index-py-aims-4"; # Add your Azure AI Search index name here


# design system prompt here
def sysPrompt(template, clientOrg=None, mspOrg=None, userRequirement=None, title=None, tech=None, 
           timeLine=None, wordLimit=None, approach=None, objectives=None,
           background=None, assumptions=None):
    prompt_template = PromptTemplate.from_template(template)
    #Template for building the PROMPT
    systemPrompt = prompt_template.format(clientOrg=clientOrg, mspOrg=mspOrg, 
                             userRequirement=userRequirement, 
                             title=title, tech=tech, timeLine=timeLine, 
                             wordLimit=wordLimit, approach=approach, 
                             objectives=objectives, background=background,
                             assumptions=assumptions)
    return systemPrompt

# Connect to your data from here
def setup_byod(deployment_id: str) -> None:
    """Sets up the OpenAI Python SDK to use your own data for the chat endpoint.

    :param deployment_id: The deployment ID for the model to use with your own data.

    To remove this configuration, simply set openai.requestssession to None.
    """

    class BringYourOwnDataAdapter(requests.adapters.HTTPAdapter):

        def send(self, request, **kwargs):
            request.url = f"{openai.api_base}/openai/deployments/{deployment_id}/extensions/chat/completions?api-version={openai.api_version}"
            return super().send(request, **kwargs)

    session = requests.Session()

    # Mount a custom adapter which will use the extensions endpoint for any call using the given `deployment_id`
    session.mount(
        prefix=f"{openai.api_base}/openai/deployments/{deployment_id}",
        adapter=BringYourOwnDataAdapter()
    )

    openai.requestssession = session
# funcion call
setup_byod(deployment_id)


# generate proposal from here
def generate_proposal():

    """
    This function helps in generating professional proposal.
    It uses azure open ai gpt-3.5-turbo model for text generation.
    text-embeddings-ada model for creating embeddings.
    And save the content as word document.
    """
    css = ''' 
    <style>
        [data-testid="stSidebar"]{
            min-width: 350px;
            max-width: 450px;
        }
        [data-testid="stSidebarContent"]{
            min-width: 300px;
            max-width: 400px;
        }
        
        [data-testid="block-container"]{
            min-width: 300px;
            max-width: 1250px;
        }
        
    </style>
    '''
    st.markdown(css, unsafe_allow_html=True)
    # Page Title
    st.title("""📄 Proposal Builder - First Draft Generator""")

    # select the proposal doc
    proposal_doc = st.sidebar.selectbox('Choose the proposal document to be generated', 
                            ('Select','RFP-Response-Template','Statement-of-Work-Template'
                             ,'Change-Order-Template', 'New-Logo-Template'),
                            index = 0, key='proposal_doc') # 'Request',
    
    # Empty section
    if proposal_doc=='Select':
        pass
    
    # section for RFP response proposal
    elif proposal_doc=='RFP-Response-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        template_doc = r'./reference_docs/template_doc.docx'
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title0')
                approach = st.text_input("What should be the scope for the doc? ", key='approach0')
                
                tech = st.text_input('What tech stack to be included?', key='tech0') # 'Request',
                clientOrg = st.text_input("What is the client organization name?", key='clientOrg0')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine0')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("Add brief about client background, Overview of Proposed Work, Current Status, Challenges Faced & Proposed Value", height=100, key='background0')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives0')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement0')
                assumptions = st.text_area("Add Basic Assumptions?", height=100, key='assumptions0')
        
        # for costs and timeline table configurations
        st.write('Add costs here')
        cost_df = pd.DataFrame(columns=['Service_Names','Quantities','MRR',
                                   'Extended_MRR'])
        cost_config = {
                    'Service_Names' : st.column_config.TextColumn('Service Names (required)', width='large', required=True),
                    'Quantities' : st.column_config.NumberColumn('Quantities', min_value=0, max_value=10),
                    'MRR' : st.column_config.NumberColumn('MRR in $'),
                    'Extended_MRR' : st.column_config.NumberColumn('Extended MRR in $')
                }
        
        costs = st.data_editor(cost_df, column_config = cost_config, num_rows='dynamic')
        costs.drop_duplicates(inplace=True, ignore_index=True)
        nd_array_costs = costs.to_dict(orient='records')
        nd_array_costs.insert(0, {'Service_Names': 'Service_Names','Quantities': 'Quantities',
                                  'MRR': 'MRR', 'Extended_MRR':'Extended_MRR'})

        st.write('Add Time Lines here')
        timeline_df = pd.DataFrame(columns=['Objectives','Weeks'])
        timeline_config = {
                    'Objectives' : st.column_config.TextColumn('Objectives', width='large', required=True),
                    'Weeks' : st.column_config.NumberColumn('Weeks', min_value=0, max_value=timeLine),
                }
        timelines = st.data_editor(timeline_df, column_config = timeline_config, num_rows='dynamic')
        # print(timelines)
        nd_array_timelines = timelines.to_dict(orient='records')
        nd_array_timelines.insert(0, {'Objectives': 'Objectives','Weeks': 'Weeks'})

        var_cost_timeline = '## COST AND TIMELINE'

        # declare response and prepare a list of it.
        # rfpResponsePromptSecTitle = sectionTitle()
        rfpResponsePromptSecIndex = sectionIndex()
        rfpResponsePromptSecA = sectionA()
        rfpResponsePromptSecB = sectionB()
        rfpResponsePromptSecC = sectionC()
        rfpResponsePromptSecD = sectionD()
        rfpResponsePromptSecE = sectionE()
        rfpResponsePromptSecF = sectionF()
        # rfpResponsePromptSecG = sectionG()
        # rfpResponsePromptSecAssump = sectionAssumptions()
        prompt_list = [rfpResponsePromptSecIndex, rfpResponsePromptSecA, rfpResponsePromptSecB, 
                                rfpResponsePromptSecC, rfpResponsePromptSecD,
                                rfpResponsePromptSecE, rfpResponsePromptSecF, 
                                ]
        
        # click submit button and call openai api
        submit = st.sidebar.button("Generate")
        if submit:
            async def callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name):
                await asyncio.sleep(10)
                completion_ = openai.ChatCompletion.create(
                    messages=message_text,
                    deployment_id=deployment_id,
                    max_tokens=8000,
                    temperature=0.5,
                    topProbabilities=0.5,
                    dataSources=[
                        {
                            "type": "AzureCognitiveSearch",
                            "parameters": {
                                "endpoint": search_endpoint,
                                "key": search_key,
                                "indexName": search_index_name,
                            }
                        }
                    ]
                )
                # Assuming there's a way to handle the completion asynchronously
                # This placeholder represents handling the response
                resp = completion_['choices'][0]['message']['content']
                return resp

            # add a streamlit progress bar
            progress_text = "Operation in progress. Please wait."
            my_bar = st.sidebar.progress(0, text=progress_text)
        
            responses_ = []
            for percent_complete, prompt in zip(range(len(prompt_list)),prompt_list):
                time.sleep(0.01)
                my_bar.progress((percent_complete + 1)*10, text=progress_text)
                sysPrompt_ = sysPrompt(prompt, mspOrg,  
                            title, timeLine, approach, tech, background, clientOrg, 
                            objectives, userRequirement, assumptions)
                message_text = [{"role": "user", "content": sysPrompt_}]
                
                # asynchronously call the openai chat api
                resp = asyncio.run(callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name))
                
                # clean response generated from openai
                if resp:
                    # patterns
                    pattern0 = r"\},\},\},\},\}\]\]"
                    pattern1 = r"\{[^}]*\}"
                    pattern2 = r"#Response|# Response|##Response|## Response|###Response|### Response"
                    pattern3 = r"#Retrieved Documents|# Retrieved Documents|##Retrieved Documents|## Retrieved Documents|###Retrieved Documents|### Retrieved Documents"
                    pattern4 = r"\[doc\d+\]"
                    # pattern5 = r"\bNote\b"
                    # pattern6 = r"\bPlease note that\b"
                    pattern5 = r"(?:Note:|Please note that|Unnecessary text and citations|For the most accurate and up-to-date information, it is recommended to refer to Synoptek's|Additional information or needs may require changes|It is concise and free from citations|However, the roles mentioned in the documents align with the requirements mentioned in the user's question.|The document is crafted in proper markdown format, reflecting the section title and subsections accurately.|The content has been crafted in proper markdown format, reflecting the section title and subsections accurately. It is clear and devoid of any external citations or documents. The content is direct and to the point, without adding unnecessary text or sentences starting with Note|This proposal is based on the RFP documentation shared between parties to date. )(.*?)(?:\.|$)"
                    pattern6 = r", and the user requirements 15.Ensure the document is crafted in proper markdown format, reflecting the section title and subsections accurately. Ensure the content is clear and devoid of any external citations or documents. Should any be present, they must be excluded. Additionally, maintain brevity and relevance in the content, avoiding the addition of unnecessary text to fulfill a word count criterion. Any found should be omitted.Additionally, the content should be direct and to the point, without adding texts which summarizes the task to meet a word count threshold. Any found should be omitted.Do not add any sentences starting with Note or ."
                    # Replace the pattern with an empty string
                    modified_resp = re.sub(pattern0, "", resp)
                    modified_resp_ = re.sub(pattern1, "", modified_resp)
                    modified_respX_ = re.sub(pattern2, "", modified_resp_)
                    modified_respX1_ = re.sub(pattern3, "", modified_respX_)
                    modified_respX2_ = re.sub(pattern4, "", modified_respX1_)
                    modified_respX3_ = re.sub(pattern5, "", modified_respX2_)
                    modified_respX4_ = re.sub(pattern6, "", modified_respX3_)
                    responses_.append(modified_respX4_)
                else:
                    print(f"No response or error for '{prompt}'")
            
            # add other texts in response list
            responses_.insert(0, title)
            assumptions_ = '\n'.join(['## ASSUMPTIONS', assumptions])
            responses_.append(assumptions_)
            # print(len(responses_))

            # declare a template with context and save it
            tpl = DocxTemplate(template_doc)
            context = {
                        "date_today": datetime.now().strftime(format='%B %d, %Y'),
                        "mspOrg":mspOrg,
                        "clientOrg":clientOrg,
                        "title": responses_[0],
                        "index": responses_[1],
                        "executive_summary": responses_[2],
                        "sow": responses_[3],
                        "tpa": responses_[4],
                        "rr": responses_[5],
                        "cr": responses_[6],
                        "tss": responses_[7],
                        "ct": var_cost_timeline,
                        "ct_table":"### COSTS",
                        "framework1": nd_array_costs,
                        "n_rows_1": costs.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_1": costs.shape[1],
                        "timeline_table": "### TIMELINES",
                        "framework2": nd_array_timelines,
                        "n_rows_2": timelines.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_2": timelines.shape[1],
                        "assumptions": responses_[8],
                        }
            jinja_env = jinja2.Environment(autoescape=True)
            tpl.render(context, jinja_env)
            tpl.save(r'./generated_docs/'+clientOrg+'_rfp_template.docx')
            
            # code for enable document download
            doc = Document(r'./generated_docs/'+clientOrg+'_rfp_template.docx')
            bio = io.BytesIO()
            doc.save(bio)
            if doc:
                st.sidebar.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name=clientOrg+'_rfp_template.docx',
                    mime="docx"
                )
            
            # additional settings for streamlit progress bar at the end of code
            time.sleep(0.1)
            my_bar.empty()
    
    # Section for Statement of Work
    elif proposal_doc=='Statement-of-Work-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        template_doc = r'./reference_docs/template_doc.docx'
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title0')
                approach = st.text_input("What should be the scope for the doc? ", key='approach0')
                
                tech = st.text_input('What tech stack to be included?', key='tech0') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg0')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine0')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("Add brief about client background, Overview of Proposed Work, Current Status, Challenges Faced & Proposed Value", height=100, key='background0')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives0')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement0')
                assumptions = st.text_area("Add Basic Assumptions?", height=100, key='assumptions0')
        
        # for costs and timeline table configurations
        st.write('Add costs here')
        cost_df = pd.DataFrame(columns=['Service_Names','Quantities','MRR',
                                   'Extended_MRR'])
        cost_config = {
                    'Service_Names' : st.column_config.TextColumn('Service Names (required)', width='large', required=True),
                    'Quantities' : st.column_config.NumberColumn('Quantities', min_value=0, max_value=10),
                    'MRR' : st.column_config.NumberColumn('MRR in $'),
                    'Extended_MRR' : st.column_config.NumberColumn('Extended MRR in $')
                }
        
        costs = st.data_editor(cost_df, column_config = cost_config, num_rows='dynamic')
        costs.drop_duplicates(inplace=True, ignore_index=True)
        nd_array_costs = costs.to_dict(orient='records')
        nd_array_costs.insert(0, {'Service_Names': 'Service_Names','Quantities': 'Quantities',
                                  'MRR': 'MRR', 'Extended_MRR':'Extended_MRR'})

        st.write('Add Time Lines here')
        timeline_df = pd.DataFrame(columns=['Objectives','Weeks'])
        timeline_config = {
                    'Objectives' : st.column_config.TextColumn('Objectives', width='large', required=True),
                    'Weeks' : st.column_config.NumberColumn('Weeks', min_value=0, max_value=timeLine),
                }
        timelines = st.data_editor(timeline_df, column_config = timeline_config, num_rows='dynamic')
        # print(timelines)
        nd_array_timelines = timelines.to_dict(orient='records')
        nd_array_timelines.insert(0, {'Objectives': 'Objectives','Weeks': 'Weeks'})

        var_cost_timeline = '## COST AND TIMELINE'

        # declare response and prepare a list of it.
        # rfpResponsePromptSecTitle = sectionTitle()
        rfpResponsePromptSecIndex = sectionIndex()
        rfpResponsePromptSecA = sectionA()
        rfpResponsePromptSecB = sectionB()
        rfpResponsePromptSecC = sectionC()
        rfpResponsePromptSecD = sectionD()
        rfpResponsePromptSecE = sectionE()
        rfpResponsePromptSecF = sectionF()
        # rfpResponsePromptSecG = sectionG()
        # rfpResponsePromptSecAssump = sectionAssumptions()
        prompt_list = [rfpResponsePromptSecIndex, rfpResponsePromptSecA, rfpResponsePromptSecB, 
                                rfpResponsePromptSecC, rfpResponsePromptSecD,
                                rfpResponsePromptSecE, rfpResponsePromptSecF, 
                                ]
        
        # click submit button and call openai api
        submit = st.sidebar.button("Generate")
        if submit:
            async def callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name):
                await asyncio.sleep(10)
                completion_ = openai.ChatCompletion.create(
                    messages=message_text,
                    deployment_id=deployment_id,
                    max_tokens=8000,
                    temperature=0.5,
                    topProbabilities=0.5,
                    dataSources=[
                        {
                            "type": "AzureCognitiveSearch",
                            "parameters": {
                                "endpoint": search_endpoint,
                                "key": search_key,
                                "indexName": search_index_name,
                            }
                        }
                    ]
                )
                # Assuming there's a way to handle the completion asynchronously
                # This placeholder represents handling the response
                resp = completion_['choices'][0]['message']['content']
                return resp

            # add a streamlit progress bar
            progress_text = "Operation in progress. Please wait."
            my_bar = st.sidebar.progress(0, text=progress_text)
        
            responses_ = []
            for percent_complete, prompt in zip(range(len(prompt_list)),prompt_list):
                time.sleep(0.01)
                my_bar.progress((percent_complete + 1)*10, text=progress_text)
                sysPrompt_ = sysPrompt(prompt, mspOrg,  
                            title, timeLine, approach, tech, background, clientOrg, 
                            objectives, userRequirement, assumptions)
                message_text = [{"role": "user", "content": sysPrompt_}]
                
                # asynchronously call the openai chat api
                resp = asyncio.run(callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name))
                
                # clean response generated from openai
                if resp:
                    # patterns
                    pattern0 = r"\},\},\},\},\}\]\]"
                    pattern1 = r"\{[^}]*\}"
                    pattern2 = r"#Response|# Response|##Response|## Response|###Response|### Response"
                    pattern3 = r"#Retrieved Documents|# Retrieved Documents|##Retrieved Documents|## Retrieved Documents|###Retrieved Documents|### Retrieved Documents"
                    pattern4 = r"\[doc\d+\]"
                    # pattern5 = r"\bNote\b"
                    # pattern6 = r"\bPlease note that\b"
                    pattern5 = r"(?:Note:|Please note that|Unnecessary text and citations|For the most accurate and up-to-date information, it is recommended to refer to Synoptek's|Additional information or needs may require changes|It is concise and free from citations|However, the roles mentioned in the documents align with the requirements mentioned in the user's question.|The document is crafted in proper markdown format, reflecting the section title and subsections accurately.|The content has been crafted in proper markdown format, reflecting the section title and subsections accurately. It is clear and devoid of any external citations or documents. The content is direct and to the point, without adding unnecessary text or sentences starting with Note|This proposal is based on the RFP documentation shared between parties to date. )(.*?)(?:\.|$)"
                    pattern6 = r", and the user requirements 15.Ensure the document is crafted in proper markdown format, reflecting the section title and subsections accurately. Ensure the content is clear and devoid of any external citations or documents. Should any be present, they must be excluded. Additionally, maintain brevity and relevance in the content, avoiding the addition of unnecessary text to fulfill a word count criterion. Any found should be omitted.Additionally, the content should be direct and to the point, without adding texts which summarizes the task to meet a word count threshold. Any found should be omitted.Do not add any sentences starting with Note or ."
                    # Replace the pattern with an empty string
                    modified_resp = re.sub(pattern0, "", resp)
                    modified_resp_ = re.sub(pattern1, "", modified_resp)
                    modified_respX_ = re.sub(pattern2, "", modified_resp_)
                    modified_respX1_ = re.sub(pattern3, "", modified_respX_)
                    modified_respX2_ = re.sub(pattern4, "", modified_respX1_)
                    modified_respX3_ = re.sub(pattern5, "", modified_respX2_)
                    modified_respX4_ = re.sub(pattern6, "", modified_respX3_)
                    responses_.append(modified_respX4_)
                else:
                    print(f"No response or error for '{prompt}'")
            
            # add other texts in response list
            responses_.insert(0, title)
            assumptions_ = '\n'.join(['## ASSUMPTIONS', assumptions])
            responses_.append(assumptions_)
            # print(len(responses_))

            # declare a template with context and save it
            tpl = DocxTemplate(template_doc)
            context = {
                        "date_today": datetime.now().strftime(format='%B %d, %Y'),
                        "mspOrg":mspOrg,
                        "clientOrg":clientOrg,
                        "title": responses_[0],
                        "index": responses_[1],
                        "executive_summary": responses_[2],
                        "sow": responses_[3],
                        "tpa": responses_[4],
                        "rr": responses_[5],
                        "cr": responses_[6],
                        "tss": responses_[7],
                        "ct": var_cost_timeline,
                        "ct_table":"### COSTS",
                        "framework1": nd_array_costs,
                        "n_rows_1": costs.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_1": costs.shape[1],
                        "timeline_table": "### TIMELINES",
                        "framework2": nd_array_timelines,
                        "n_rows_2": timelines.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_2": timelines.shape[1],
                        "assumptions": responses_[8],
                        }
            jinja_env = jinja2.Environment(autoescape=True)
            tpl.render(context, jinja_env)
            tpl.save(r'./generated_docs/'+clientOrg+'_sow_template.docx')
            
            # code for enable document download
            doc = Document(r'./generated_docs/'+clientOrg+'_sow_template.docx')
            bio = io.BytesIO()
            doc.save(bio)
            if doc:
                st.sidebar.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name=clientOrg+'_sow_template.docx',
                    mime="docx"
                )
            
            # additional settings for streamlit progress bar at the end of code
            time.sleep(0.1)
            my_bar.empty()
    
    # Section for Change Order
    elif proposal_doc=='Change-Order-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        template_doc = r'./reference_docs/template_doc.docx'
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title0')
                approach = st.text_input("What should be the scope for the doc? ", key='approach0')
                
                tech = st.text_input('What tech stack to be included?', key='tech0') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg0')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine0')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("Add brief about client background, Overview of Proposed Work, Current Status, Challenges Faced & Proposed Value", height=100, key='background0')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives0')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement0')
                assumptions = st.text_area("Add Basic Assumptions?", height=100, key='assumptions0')
        
        # for costs and timeline table configurations
        st.write('Add costs here')
        cost_df = pd.DataFrame(columns=['Service_Names','Quantities','MRR',
                                   'Extended_MRR'])
        cost_config = {
                    'Service_Names' : st.column_config.TextColumn('Service Names (required)', width='large', required=True),
                    'Quantities' : st.column_config.NumberColumn('Quantities', min_value=0, max_value=10),
                    'MRR' : st.column_config.NumberColumn('MRR in $'),
                    'Extended_MRR' : st.column_config.NumberColumn('Extended MRR in $')
                }
        
        costs = st.data_editor(cost_df, column_config = cost_config, num_rows='dynamic')
        costs.drop_duplicates(inplace=True, ignore_index=True)
        nd_array_costs = costs.to_dict(orient='records')
        nd_array_costs.insert(0, {'Service_Names': 'Service_Names','Quantities': 'Quantities',
                                  'MRR': 'MRR', 'Extended_MRR':'Extended_MRR'})

        st.write('Add Time Lines here')
        timeline_df = pd.DataFrame(columns=['Objectives','Weeks'])
        timeline_config = {
                    'Objectives' : st.column_config.TextColumn('Objectives', width='large', required=True),
                    'Weeks' : st.column_config.NumberColumn('Weeks', min_value=0, max_value=timeLine),
                }
        timelines = st.data_editor(timeline_df, column_config = timeline_config, num_rows='dynamic')
        # print(timelines)
        nd_array_timelines = timelines.to_dict(orient='records')
        nd_array_timelines.insert(0, {'Objectives': 'Objectives','Weeks': 'Weeks'})

        var_cost_timeline = '## COST AND TIMELINE'

        # declare response and prepare a list of it.
        # rfpResponsePromptSecTitle = sectionTitle()
        rfpResponsePromptSecIndex = sectionIndex()
        rfpResponsePromptSecA = sectionA()
        rfpResponsePromptSecB = sectionB()
        rfpResponsePromptSecC = sectionC()
        rfpResponsePromptSecD = sectionD()
        rfpResponsePromptSecE = sectionE()
        rfpResponsePromptSecF = sectionF()
        # rfpResponsePromptSecG = sectionG()
        # rfpResponsePromptSecAssump = sectionAssumptions()
        prompt_list = [rfpResponsePromptSecIndex, rfpResponsePromptSecA, rfpResponsePromptSecB, 
                                rfpResponsePromptSecC, rfpResponsePromptSecD,
                                rfpResponsePromptSecE, rfpResponsePromptSecF, 
                                ]
        
        # click submit button and call openai api
        submit = st.sidebar.button("Generate")
        if submit:
            async def callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name):
                await asyncio.sleep(10)
                completion_ = openai.ChatCompletion.create(
                    messages=message_text,
                    deployment_id=deployment_id,
                    max_tokens=8000,
                    temperature=0.5,
                    topProbabilities=0.5,
                    dataSources=[
                        {
                            "type": "AzureCognitiveSearch",
                            "parameters": {
                                "endpoint": search_endpoint,
                                "key": search_key,
                                "indexName": search_index_name,
                            }
                        }
                    ]
                )
                # Assuming there's a way to handle the completion asynchronously
                # This placeholder represents handling the response
                resp = completion_['choices'][0]['message']['content']
                return resp

            # add a streamlit progress bar
            progress_text = "Operation in progress. Please wait."
            my_bar = st.sidebar.progress(0, text=progress_text)
        
            responses_ = []
            for percent_complete, prompt in zip(range(len(prompt_list)),prompt_list):
                time.sleep(0.01)
                my_bar.progress((percent_complete + 1)*10, text=progress_text)
                sysPrompt_ = sysPrompt(prompt, mspOrg,  
                            title, timeLine, approach, tech, background, clientOrg, 
                            objectives, userRequirement, assumptions)
                message_text = [{"role": "user", "content": sysPrompt_}]
                
                # asynchronously call the openai chat api
                resp = asyncio.run(callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name))
                
                # clean response generated from openai
                if resp:
                    # patterns
                    pattern0 = r"\},\},\},\},\}\]\]"
                    pattern1 = r"\{[^}]*\}"
                    pattern2 = r"#Response|# Response|##Response|## Response|###Response|### Response"
                    pattern3 = r"#Retrieved Documents|# Retrieved Documents|##Retrieved Documents|## Retrieved Documents|###Retrieved Documents|### Retrieved Documents"
                    pattern4 = r"\[doc\d+\]"
                    # pattern5 = r"\bNote\b"
                    # pattern6 = r"\bPlease note that\b"
                    pattern5 = r"(?:Note:|Please note that|Unnecessary text and citations|For the most accurate and up-to-date information, it is recommended to refer to Synoptek's|Additional information or needs may require changes|It is concise and free from citations|However, the roles mentioned in the documents align with the requirements mentioned in the user's question.|The document is crafted in proper markdown format, reflecting the section title and subsections accurately.|The content has been crafted in proper markdown format, reflecting the section title and subsections accurately. It is clear and devoid of any external citations or documents. The content is direct and to the point, without adding unnecessary text or sentences starting with Note|This proposal is based on the RFP documentation shared between parties to date. )(.*?)(?:\.|$)"
                    pattern6 = r", and the user requirements 15.Ensure the document is crafted in proper markdown format, reflecting the section title and subsections accurately. Ensure the content is clear and devoid of any external citations or documents. Should any be present, they must be excluded. Additionally, maintain brevity and relevance in the content, avoiding the addition of unnecessary text to fulfill a word count criterion. Any found should be omitted.Additionally, the content should be direct and to the point, without adding texts which summarizes the task to meet a word count threshold. Any found should be omitted.Do not add any sentences starting with Note or ."
                    # Replace the pattern with an empty string
                    modified_resp = re.sub(pattern0, "", resp)
                    modified_resp_ = re.sub(pattern1, "", modified_resp)
                    modified_respX_ = re.sub(pattern2, "", modified_resp_)
                    modified_respX1_ = re.sub(pattern3, "", modified_respX_)
                    modified_respX2_ = re.sub(pattern4, "", modified_respX1_)
                    modified_respX3_ = re.sub(pattern5, "", modified_respX2_)
                    modified_respX4_ = re.sub(pattern6, "", modified_respX3_)
                    responses_.append(modified_respX4_)
                else:
                    print(f"No response or error for '{prompt}'")
            
            # add other texts in response list
            responses_.insert(0, title)
            assumptions_ = '\n'.join(['## ASSUMPTIONS', assumptions])
            responses_.append(assumptions_)
            # print(len(responses_))

            # declare a template with context and save it
            tpl = DocxTemplate(template_doc)
            context = {
                        "date_today": datetime.now().strftime(format='%B %d, %Y'),
                        "mspOrg":mspOrg,
                        "clientOrg":clientOrg,
                        "title": responses_[0],
                        "index": responses_[1],
                        "executive_summary": responses_[2],
                        "sow": responses_[3],
                        "tpa": responses_[4],
                        "rr": responses_[5],
                        "cr": responses_[6],
                        "tss": responses_[7],
                        "ct": var_cost_timeline,
                        "ct_table":"### COSTS",
                        "framework1": nd_array_costs,
                        "n_rows_1": costs.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_1": costs.shape[1],
                        "timeline_table": "### TIMELINES",
                        "framework2": nd_array_timelines,
                        "n_rows_2": timelines.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_2": timelines.shape[1],
                        "assumptions": responses_[8],
                        }
            jinja_env = jinja2.Environment(autoescape=True)
            tpl.render(context, jinja_env)
            tpl.save(r'./generated_docs/'+clientOrg+'_co_template.docx')
            
            # code for enable document download
            doc = Document(r'./generated_docs/'+clientOrg+'_co_template.docx')
            bio = io.BytesIO()
            doc.save(bio)
            if doc:
                st.sidebar.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name=clientOrg+'_co_template.docx',
                    mime="docx"
                )
            
            # additional settings for streamlit progress bar at the end of code
            time.sleep(0.1)
            my_bar.empty()
    
    # Section for New-Logo-Template
    elif proposal_doc=='New-Logo-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        template_doc = r'./reference_docs/template_doc.docx'
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title0')
                approach = st.text_input("What should be the scope for the doc? ", key='approach0')
                
                tech = st.text_input('What tech stack to be included?', key='tech0') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg0')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine0')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("Add brief about client background, Overview of Proposed Work, Current Status, Challenges Faced & Proposed Value", height=100, key='background0')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives0')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement0')
                assumptions = st.text_area("Add Basic Assumptions?", height=100, key='assumptions0')
        
        # for costs and timeline table configurations
        st.write('Add costs here')
        cost_df = pd.DataFrame(columns=['Service_Names','Quantities','MRR',
                                   'Extended_MRR'])
        cost_config = {
                    'Service_Names' : st.column_config.TextColumn('Service Names (required)', width='large', required=True),
                    'Quantities' : st.column_config.NumberColumn('Quantities', min_value=0, max_value=10),
                    'MRR' : st.column_config.NumberColumn('MRR in $'),
                    'Extended_MRR' : st.column_config.NumberColumn('Extended MRR in $')
                }
        
        costs = st.data_editor(cost_df, column_config = cost_config, num_rows='dynamic')
        costs.drop_duplicates(inplace=True, ignore_index=True)
        nd_array_costs = costs.to_dict(orient='records')
        nd_array_costs.insert(0, {'Service_Names': 'Service_Names','Quantities': 'Quantities',
                                  'MRR': 'MRR', 'Extended_MRR':'Extended_MRR'})

        st.write('Add Time Lines here')
        timeline_df = pd.DataFrame(columns=['Objectives','Weeks'])
        timeline_config = {
                    'Objectives' : st.column_config.TextColumn('Objectives', width='large', required=True),
                    'Weeks' : st.column_config.NumberColumn('Weeks', min_value=0, max_value=timeLine),
                }
        timelines = st.data_editor(timeline_df, column_config = timeline_config, num_rows='dynamic')
        # print(timelines)
        nd_array_timelines = timelines.to_dict(orient='records')
        nd_array_timelines.insert(0, {'Objectives': 'Objectives','Weeks': 'Weeks'})

        var_cost_timeline = '## COST AND TIMELINE'

        # declare response and prepare a list of it.
        # rfpResponsePromptSecTitle = sectionTitle()
        rfpResponsePromptSecIndex = sectionIndex()
        rfpResponsePromptSecA = sectionA()
        rfpResponsePromptSecB = sectionB()
        rfpResponsePromptSecC = sectionC()
        rfpResponsePromptSecD = sectionD()
        rfpResponsePromptSecE = sectionE()
        rfpResponsePromptSecF = sectionF()
        # rfpResponsePromptSecG = sectionG()
        # rfpResponsePromptSecAssump = sectionAssumptions()
        prompt_list = [rfpResponsePromptSecIndex, rfpResponsePromptSecA, rfpResponsePromptSecB, 
                                rfpResponsePromptSecC, rfpResponsePromptSecD,
                                rfpResponsePromptSecE, rfpResponsePromptSecF, 
                                ]
        
        # click submit button and call openai api
        submit = st.sidebar.button("Generate")
        if submit:
            async def callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name):
                await asyncio.sleep(10)
                completion_ = openai.ChatCompletion.create(
                    messages=message_text,
                    deployment_id=deployment_id,
                    max_tokens=8000,
                    temperature=0.5,
                    topProbabilities=0.5,
                    dataSources=[
                        {
                            "type": "AzureCognitiveSearch",
                            "parameters": {
                                "endpoint": search_endpoint,
                                "key": search_key,
                                "indexName": search_index_name,
                            }
                        }
                    ]
                )
                # Assuming there's a way to handle the completion asynchronously
                # This placeholder represents handling the response
                resp = completion_['choices'][0]['message']['content']
                return resp

            # add a streamlit progress bar
            progress_text = "Operation in progress. Please wait."
            my_bar = st.sidebar.progress(0, text=progress_text)
        
            responses_ = []
            for percent_complete, prompt in zip(range(len(prompt_list)),prompt_list):
                time.sleep(0.01)
                my_bar.progress((percent_complete + 1)*10, text=progress_text)
                sysPrompt_ = sysPrompt(prompt, mspOrg,  
                            title, timeLine, approach, tech, background, clientOrg, 
                            objectives, userRequirement, assumptions)
                message_text = [{"role": "user", "content": sysPrompt_}]
                
                # asynchronously call the openai chat api
                resp = asyncio.run(callOpenAI(message_text, deployment_id, search_endpoint, search_key, search_index_name))
                
                # clean response generated from openai
                if resp:
                    # patterns
                    pattern0 = r"\},\},\},\},\}\]\]"
                    pattern1 = r"\{[^}]*\}"
                    pattern2 = r"#Response|# Response|##Response|## Response|###Response|### Response"
                    pattern3 = r"#Retrieved Documents|# Retrieved Documents|##Retrieved Documents|## Retrieved Documents|###Retrieved Documents|### Retrieved Documents"
                    pattern4 = r"\[doc\d+\]"
                    # pattern5 = r"\bNote\b"
                    # pattern6 = r"\bPlease note that\b"
                    pattern5 = r"(?:Note:|Please note that|Unnecessary text and citations|For the most accurate and up-to-date information, it is recommended to refer to Synoptek's|Additional information or needs may require changes|It is concise and free from citations|However, the roles mentioned in the documents align with the requirements mentioned in the user's question.|The document is crafted in proper markdown format, reflecting the section title and subsections accurately.|The content has been crafted in proper markdown format, reflecting the section title and subsections accurately. It is clear and devoid of any external citations or documents. The content is direct and to the point, without adding unnecessary text or sentences starting with Note|This proposal is based on the RFP documentation shared between parties to date. )(.*?)(?:\.|$)"
                    pattern6 = r", and the user requirements 15.Ensure the document is crafted in proper markdown format, reflecting the section title and subsections accurately. Ensure the content is clear and devoid of any external citations or documents. Should any be present, they must be excluded. Additionally, maintain brevity and relevance in the content, avoiding the addition of unnecessary text to fulfill a word count criterion. Any found should be omitted.Additionally, the content should be direct and to the point, without adding texts which summarizes the task to meet a word count threshold. Any found should be omitted.Do not add any sentences starting with Note or ."
                    # Replace the pattern with an empty string
                    modified_resp = re.sub(pattern0, "", resp)
                    modified_resp_ = re.sub(pattern1, "", modified_resp)
                    modified_respX_ = re.sub(pattern2, "", modified_resp_)
                    modified_respX1_ = re.sub(pattern3, "", modified_respX_)
                    modified_respX2_ = re.sub(pattern4, "", modified_respX1_)
                    modified_respX3_ = re.sub(pattern5, "", modified_respX2_)
                    modified_respX4_ = re.sub(pattern6, "", modified_respX3_)
                    responses_.append(modified_respX4_)
                else:
                    print(f"No response or error for '{prompt}'")
            
            # add other texts in response list
            responses_.insert(0, title)
            assumptions_ = '\n'.join(['## ASSUMPTIONS', assumptions])
            responses_.append(assumptions_)
            # print(len(responses_))

            # declare a template with context and save it
            tpl = DocxTemplate(template_doc)
            context = {
                        "date_today": datetime.now().strftime(format='%B %d, %Y'),
                        "mspOrg":mspOrg,
                        "clientOrg":clientOrg,
                        "title": responses_[0],
                        "index": responses_[1],
                        "executive_summary": responses_[2],
                        "sow": responses_[3],
                        "tpa": responses_[4],
                        "rr": responses_[5],
                        "cr": responses_[6],
                        "tss": responses_[7],
                        "ct": var_cost_timeline,
                        "ct_table":"### COSTS",
                        "framework1": nd_array_costs,
                        "n_rows_1": costs.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_1": costs.shape[1],
                        "timeline_table": "### TIMELINES",
                        "framework2": nd_array_timelines,
                        "n_rows_2": timelines.shape[0]-1,  # note: minus 1 here because we added an extra row for the headers
                        "n_columns_2": timelines.shape[1],
                        "assumptions": responses_[8],
                        }
            jinja_env = jinja2.Environment(autoescape=True)
            tpl.render(context, jinja_env)
            tpl.save(r'./generated_docs/'+clientOrg+'_nl_template.docx')
            
            # code for enable document download
            doc = Document(r'./generated_docs/'+clientOrg+'_nl_template.docx')
            bio = io.BytesIO()
            doc.save(bio)
            if doc:
                st.sidebar.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name=clientOrg+'_nl_template.docx',
                    mime="docx"
                )
            
            # additional settings for streamlit progress bar at the end of code
            time.sleep(0.1)
            my_bar.empty()