"""
This code uses python==3.11.4, and open_ai==0.28.1, langchain==0.0.354
"""
# load packages
import streamlit as st
from streamlit_quill import st_quill
from langchain.prompts import PromptTemplate
from langchain.llms import CTransformers, OpenAI
import openai
import pandas
import requests
import pypandoc
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from azure.storage.blob import BlobServiceClient

import os
import io
from dotenv import load_dotenv,find_dotenv
load_dotenv(find_dotenv())

from prompts.allDoc_Prompts import *
from azure.identity import DefaultAzureCredential, ManagedIdentityCredential
from azure.keyvault.secrets import SecretClient

# # azure key-vault authentication
# CREDENTIAL = DefaultAzureCredential()
# client = SecretClient(
#     vault_url="https://kayvault-aims.vault.azure.net/",
#     credential=CREDENTIAL
# )
# print(client.get_secret('azure-openai-endpoint'))

# local folder setup
if not os.path.exists(r'./generated_docs/'):
    os.mkdir(r'./generated_docs/')

# azure blob storage setup
# container_name = "container-aims-output"

# def uploadToBlobStorage(file_path,file_name):
#     blob_service_client = BlobServiceClient.from_connection_string(os.getenv("AZURE_BLOB_CONNECTION_STRING"))
#     blob_client = blob_service_client.get_blob_client(container=container_name, blob=file_name)
#     with open(file_path,"rb") as data:
#         blob_client.upload_blob(data)
#     print(f"Uploaded {file_name}.")

# azure open ai setup
openai.api_type = "azure"
openai.api_base = "https://dev-aims-ai.openai.azure.com/"
openai.api_version = "2023-12-01-preview"
openai.api_key = os.getenv("OPENAI_API_KEY_AZURE")# "4921760f8fd2477d894fce4c2e4e3e54"# os.getenv("OPENAI_API_KEY_AZURE")

deployment_id = "gpt-35-aims" # Add your deployment ID here

# Azure AI Search setup
search_endpoint = "https://cogs-aims.search.windows.net"; # Add your Azure AI Search endpoint here
search_key = os.getenv("SEARCH_KEY")# "N1YwVm7QvGpRydRWf403VewtIDr5blIxgZqdRVh6fDAzSeBIhN6O"# os.getenv("SEARCH_KEY"); # Add your Azure AI Search admin key here
search_index_name = "try-index-py-aims-4"; # Add your Azure AI Search index name here

# design prompt here
def prompt(template, clientOrg=None, mspOrg=None, userRequirement=None, title=None, tech=None, 
           timeLine=None, approach=None, objectives=None,
           background=None):
    #Template for building the PROMPT
    prompt = template.format(clientOrg=clientOrg, mspOrg=mspOrg, 
                             userRequirement=userRequirement, 
                             title=title, tech=tech, timeLine=timeLine, 
                             approach=approach, 
                             objectives=objectives, background=background)
    return prompt

# Connect to your data from here
def setup_byod(deployment_id: str) -> None:
    """Sets up the OpenAI Python SDK to use your own data for the chat endpoint.

    :param deployment_id: The deployment ID for the model to use with your own data.

    To remove this configuration, simply set openai.requestssession to None.
    """

    class BringYourOwnDataAdapter(requests.adapters.HTTPAdapter):

        def send(self, request, **kwargs):
            request.url = f"{openai.api_base}/openai/deployments/{deployment_id}/extensions/chat/completions?api-version={openai.api_version}"
            return super().send(request, **kwargs)

    session = requests.Session()

    # Mount a custom adapter which will use the extensions endpoint for any call using the given `deployment_id`
    session.mount(
        prefix=f"{openai.api_base}/openai/deployments/{deployment_id}",
        adapter=BringYourOwnDataAdapter()
    )

    openai.requestssession = session
# funcion call
setup_byod(deployment_id)


# generate proposal from here
def generate_proposal():

    """
    This function helps in generating professional proposal.
    It uses azure open ai gpt-3.5-turbo model for text generation.
    text-embeddings-ada model for creating embeddings.
    And save the content as word document.
    """
    css = ''' 
    <style>
        [data-testid="stSidebar"]{
            min-width: 350px;
            max-width: 450px;
        }
        [data-testid="stSidebarContent"]{
            min-width: 300px;
            max-width: 400px;
        }
        
        [data-testid="block-container"]{
            min-width: 300px;
            max-width: 1250px;
        }
        
    </style>
    '''
    st.markdown(css, unsafe_allow_html=True)

    
    
    # Page Title
    st.title("""📄 Proposal Builder - First Draft Generator""")

    # select the proposal doc
    proposal_doc = st.sidebar.selectbox('Choose the proposal document to be generated', 
                            ('Select','RFP-Response-Template','Statement-of-Work-Template'
                             ,'Change-Order-Template', 'New-Logo-Template'),
                            index = 0, key='proposal_doc') # 'Request',
    
    # Empty section
    if proposal_doc=='Select':
        pass
    
    # section for RFP response proposal
    elif proposal_doc=='RFP-Response-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        reference_docx = r'./reference_docs/reference.docx'
        #template_doc = r'./reference_docs/template_doc.docx'
    
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title0')
                approach = st.text_input("What should be the scope for the doc? ", key='approach0')
                
                tech = st.text_input('What tech stack to be included?', key='tech0') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg0')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine0')
                # wordLimit = st.text_input("How much word count needed for proposal?", key='wordLimit0')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("What is the client background?", height=100, key='background0')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives0')
                # sections = st.multiselect('What sections to be included?',
                #                           ['EXECUTIVE SUMMARY', 'SCOPE OF WORK', 'INTEGRATED SOLUTION APPROACH', 
                #                            'TECH STACK & SERVICES','TRANSITION PLAN AND APPROACH','ROLES AND RESPONSIBILITIES',
                #                            'TERMS AND CONDITIONS, ASSUMPTIONS','COST AND TIMELINE','DELIVERABLES',
                #                            'APPROVAL','APPENDIX – RFP FAQs'],key='sections0')
                
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement0')
                submit = st.button("Generate")
                rfpResponsePrompt = rfp_response_Main()
                prmpt = prompt(rfpResponsePrompt, clientOrg, mspOrg, userRequirement, 
                               title, tech, timeLine, approach)

                message_text = [{"role": "user", "content": prmpt}]
                
                

                #When 'Generate' button is clicked, execute the below code
                
                if submit:
                    with st.spinner("Generating Response ... "):
                        
                        completion = openai.ChatCompletion.create(
                            messages=message_text,
                            deployment_id=deployment_id,
                            max_tokens = 6000,
                            temperature=0.7,
                            topProbablities=0.3,
                            frequencyPenalty=0,
                            presencePenalty=0,
                            dataSources=[  # camelCase is intentional, as this is the format the API expects
                                {
                                    "type": "AzureCognitiveSearch",
                                    "parameters": {
                                        "endpoint": search_endpoint,
                                        "key": search_key,
                                        "indexName": search_index_name,
                                    }
                                }
                            ]
                        )
                        # Update session state with the new response
                        resp = completion['choices'][0]['message']['content']
                        
                        output = pypandoc.convert_text(resp, 
                                               'docx', 
                                               format='markdown+multiline_tables', 
                                               outputfile=r'./generated_docs/'+clientOrg+'_rfp_template.docx', 
                                               extra_args=['--reference-doc=' + reference_docx, 
                                                           '-s'])
                        doc = Document(r'./generated_docs/'+clientOrg+'_rfp_template.docx')
                        bio = io.BytesIO()
                        doc.save(bio)
                        if doc:
                            st.download_button(
                                label="Click here to download",
                                data=bio.getvalue(),
                                file_name=clientOrg+'_rfp_template.docx',
                                mime="docx"
                            )
                

    # Section for Statement of Work
    elif proposal_doc=='Statement-of-Work-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        reference_docx = r'./reference_docs/reference.docx'
    
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title1')
                approach = st.text_input("What should be the scope for the doc? ", key='approach1')
                tech = st.text_input('What tech stack to be included?', key='tech1') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg1')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine1')
                wordLimit = st.text_input("How much word count needed for proposal?", key='wordLimit1')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("What is the client background?", height=100, key='background1')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives1')
                # sections = st.multiselect('What sections to be included?',
                #                           ['EXECUTIVE SUMMARY', 'SCOPE OF WORK', 'INTEGRATED SOLUTION APPROACH', 
                #                            'TECH STACK & SERVICES','TRANSITION PLAN AND APPROACH','ROLES AND RESPONSIBILITIES',
                #                            'TERMS AND CONDITIONS, ASSUMPTIONS','COST AND TIMELINE','DELIVERABLES',
                #                            'APPROVAL','APPENDIX – RFP FAQs'],key='sections1')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement1')
                submit = st.button("Generate")
                rfpResponsePrompt = statement_of_work_Main()
                prmpt = prompt(rfpResponsePrompt, clientOrg, mspOrg, userRequirement, 
                               title, tech, timeLine, approach, objectives, background)

                message_text = [{"role": "user", "content": prmpt}]

                #When 'Generate' button is clicked, execute the below code
                if submit:
                    
                    with st.spinner("Generating Response ... "):
                        
                        completion = openai.ChatCompletion.create(
                            messages=message_text,
                            deployment_id=deployment_id,
                            max_tokens = 6000,
                            temperature=0.7,
                            topProbablities=0.3,
                            frequencyPenalty=0,
                            presencePenalty=0,
                            dataSources=[  # camelCase is intentional, as this is the format the API expects
                                {
                                    "type": "AzureCognitiveSearch",
                                    "parameters": {
                                        "endpoint": search_endpoint,
                                        "key": search_key,
                                        "indexName": search_index_name,
                                    }
                                }
                            ]
                        )
                        resp = completion['choices'][0]['message']['content']
                        output = pypandoc.convert_text(resp, 
                                               'docx', 
                                               format='markdown+multiline_tables', 
                                               outputfile=r'./generated_docs/'+clientOrg+'_sow_template.docx', 
                                               extra_args=['--reference-doc=' + reference_docx, 
                                                           '-s'])
                        doc = Document(r'./generated_docs/'+clientOrg+'_sow_template.docx')
                        bio = io.BytesIO()
                        doc.save(bio)
                        if doc:
                            st.download_button(
                                label="Click here to download",
                                data=bio.getvalue(),
                                file_name=clientOrg+'_sow_template.docx',
                                mime="docx"
                            )
    # Section for Change Order
    elif proposal_doc=='Change-Order-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        reference_docx = r'./reference_docs/reference.docx'
    
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title2')
                approach = st.text_input("What should be the scope for the doc? ", key='approach2')
                tech = st.text_input('What tech stack to be included?', key='tech2') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg2')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine2')
                wordLimit = st.text_input("How much word count needed for proposal?", key='wordLimit2')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("What is the client background?", height=100, key='background2')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives2')
                # sections = st.multiselect('What sections to be included?',
                #                           ['EXECUTIVE SUMMARY', 'SCOPE OF WORK', 'INTEGRATED SOLUTION APPROACH', 
                #                            'TECH STACK & SERVICES','TRANSITION PLAN AND APPROACH','ROLES AND RESPONSIBILITIES',
                #                            'TERMS AND CONDITIONS, ASSUMPTIONS','COST AND TIMELINE','DELIVERABLES',
                #                            'APPROVAL','APPENDIX – RFP FAQs'],key='sections2')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement2')
                submit = st.button("Generate")
                rfpResponsePrompt = change_Order_Main()
                prmpt = prompt(rfpResponsePrompt, clientOrg, mspOrg, userRequirement, 
                               title, tech, timeLine, approach)

                message_text = [{"role": "user", "content": prmpt}]

                #When 'Generate' button is clicked, execute the below code
                if submit:
                    
                    with st.spinner("Generating Response ... "):
                        
                        completion = openai.ChatCompletion.create(
                            messages=message_text,
                            deployment_id=deployment_id,
                            max_tokens = 6000,
                            temperature=0.7,
                            topProbablities=0.3,
                            frequencyPenalty=0,
                            presencePenalty=0,
                            dataSources=[  # camelCase is intentional, as this is the format the API expects
                                {
                                    "type": "AzureCognitiveSearch",
                                    "parameters": {
                                        "endpoint": search_endpoint,
                                        "key": search_key,
                                        "indexName": search_index_name,
                                    }
                                }
                            ]
                        )
                        resp = completion['choices'][0]['message']['content']
                        output = pypandoc.convert_text(resp, 
                                               'docx', 
                                               format='markdown+multiline_tables', 
                                               outputfile=r'./generated_docs/'+clientOrg+'_co_template.docx', 
                                               extra_args=['--reference-doc=' + reference_docx, 
                                                           '-s'])
                        doc = Document(r'./generated_docs/'+clientOrg+'_co_template.docx')
                        bio = io.BytesIO()
                        doc.save(bio)
                        if doc:
                            st.download_button(
                                label="Click here to download",
                                data=bio.getvalue(),
                                file_name=clientOrg+'_co_template.docx',
                                mime="docx"
                            )
    # Section for New-Logo-Template
    elif proposal_doc=='New-Logo-Template':
        # Path to your custom reference.docx with "Segoe UI" font
        reference_docx = r'./reference_docs/reference.docx'
    
        col1, col2 = st.columns([1,3])
        with col1:
            with st.container(border=True):
                title = st.text_input("What should be the title for the doc? ", key='title3')
                approach = st.text_input("What should be the scope for the doc? ", key='approach3')
                tech = st.text_input('What tech stack to be included?', key='tech3') # 'Request',
                clientOrg = st.text_input("What is the client orgnisation name?", key='clientOrg3')
                timeLine = st.text_input("How much is the tentative project duration in weeks?",key='timeLine3')
                wordLimit = st.text_input("How much word count needed for proposal?", key='wordLimit3')
                mspOrg = "SYNOPTEK"

        with col2:
            with st.container(border=True):
                background = st.text_area("What is the client background?", height=100, key='background3')
                objectives = st.text_area("What are the objectives?", height=100, key='objectives3')
                # sections = st.multiselect('What sections to be included?',
                #                           ['EXECUTIVE SUMMARY', 'SCOPE OF WORK', 'INTEGRATED SOLUTION APPROACH', 
                #                            'TECH STACK & SERVICES','TRANSITION PLAN AND APPROACH','ROLES AND RESPONSIBILITIES',
                #                            'TERMS AND CONDITIONS, ASSUMPTIONS','COST AND TIMELINE','DELIVERABLES',
                #                            'APPROVAL','APPENDIX – RFP FAQs'],key='sections3')
                userRequirement = st.text_area("What are the user requirement?", height=100, key='userRequirement3')
                submit = st.button("Generate")
                rfpResponsePrompt = new_Logo_Main()
                prmpt = prompt(rfpResponsePrompt, clientOrg, mspOrg, userRequirement, 
                               title, tech, timeLine, approach)

                message_text = [{"role": "user", "content": prmpt}]

                #When 'Generate' button is clicked, execute the below code
                if submit:
                    
                    with st.spinner("Generating Response ... "):
                        
                        completion = openai.ChatCompletion.create(
                            messages=message_text,
                            deployment_id=deployment_id,
                            max_tokens = 6200,
                            temperature=0.1,
                            topProbablities=0.9,
                            frequencyPenalty=0.2,
                            presencePenalty=0.2,
                            dataSources=[  # camelCase is intentional, as this is the format the API expects
                                {
                                    "type": "AzureCognitiveSearch",
                                    "parameters": {
                                        "endpoint": search_endpoint,
                                        "key": search_key,
                                        "indexName": search_index_name,
                                    }
                                }
                            ]
                        )
                        resp = completion['choices'][0]['message']['content']
                        output = pypandoc.convert_text(resp, 
                                               'docx', 
                                               format='markdown+multiline_tables', 
                                               outputfile=r'./generated_docs/'+clientOrg+'_nl_template.docx', 
                                               extra_args=['--reference-doc=' + reference_docx, 
                                                           '-s'])
                        doc = Document(r'./generated_docs/'+clientOrg+'_nl_template.docx')
                        bio = io.BytesIO()
                        doc.save(bio)
                        if doc:
                            st.download_button(
                                label="Click here to download",
                                data=bio.getvalue(),
                                file_name=clientOrg+'_nl_template.docx',
                                mime="docx"
                            )